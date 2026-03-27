"""
01_ingestion.py — 文档摄入与切割

学习重点：
  1. 多格式文档解析（TXT / PDF / Word）
  2. 两种 Chunking 策略的原理与取舍
  3. 元数据（Metadata）设计 —— 后续过滤、引用溯源的基础
  4. chunk_overlap 为什么必须有
"""

import os
import re
import hashlib
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import numpy as np


# ─────────────────────────────────────────
# 数据结构
# ─────────────────────────────────────────

@dataclass
class Document:
    """
    原始文档：从磁盘/系统读进来的原始内容。
    还没切割，保留完整文本。
    """
    content: str
    metadata: dict = field(default_factory=dict)
    # metadata 示例：
    # {
    #   "source": "hr_policy.txt",
    #   "doc_type": "policy",          # 用于检索时按类型过滤
    #   "department": "HR",
    #   "version": "2024",
    #   "content_hash": "abc123...",   # 增量更新时判断文档是否变化
    # }


@dataclass
class Chunk:
    """
    切割后的文本块：进入向量库的最小单元。
    注意：chunk 的 metadata 比 document 更细粒度（多了页码、chunk序号等）
    """
    content: str
    metadata: dict = field(default_factory=dict)
    # metadata 示例：
    # {
    #   "source": "hr_policy.txt",
    #   "chunk_index": 3,              # 这是文档的第几个 chunk
    #   "chunk_total": 12,             # 文档共切了多少个 chunk
    #   "start_char": 450,             # 在原文中的起始字符位置（用于溯源）
    #   "section": "第二章 假期管理制度", # 所属章节（如果能提取到）
    #   "doc_type": "policy",
    #   "content_hash": "abc123...",   # chunk 内容的 hash，增量更新用
    # }


# ─────────────────────────────────────────
# 第一步：文档解析
# ─────────────────────────────────────────

class DocumentParser:
    """
    负责把各种格式的文件读成统一的 Document 对象。

    技术要点：
    - 不同格式用不同库，但统一输出 Document
    - 解析时就提取尽可能多的元数据（后面很难补）
    - content_hash 在这里生成，用于增量更新时判断文档是否修改过
    """

    def parse(self, file_path: str) -> Document:
        path = Path(file_path)
        ext = path.suffix.lower()

        parsers = {
            ".txt": self._parse_txt,
            ".md":  self._parse_txt,    # Markdown 当纯文本处理（保留原始格式）
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
        }

        if ext not in parsers:
            raise ValueError(f"不支持的文件格式: {ext}")

        content = parsers[ext](str(path))

        # 基础元数据：所有格式都有的
        metadata = {
            "source":       path.name,
            "source_path":  str(path.absolute()),
            "doc_type":     self._infer_doc_type(path.name),
            "file_size":    path.stat().st_size,
            "content_hash": hashlib.md5(content.encode()).hexdigest(),
            # ↑ MD5 够用，不需要 SHA256。用途是"判断是否变化"，不是安全场景
        }

        return Document(content=content, metadata=metadata)

    def _parse_txt(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    def _parse_pdf(self, file_path: str) -> str:
        """
        PDF 解析的几个坑：
        1. 双栏布局：fitz 默认按坐标顺序读，双栏会左右交错。
           解决：用 sort=True 参数，或按列分块再合并。
        2. 扫描版 PDF：图片 PDF，fitz 读不到文字，需要 OCR。
           判断方法：如果提取文字量极少（< 100 字/页），认为是扫描版。
        3. 页眉页脚噪声：每页顶部/底部的"第X页""公司名称"会污染 chunk。
           这里用简单的规则过滤，生产环境可以用坐标过滤（y < 50 或 y > 750）。
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("请安装 PyMuPDF: pip install PyMuPDF")

        doc = fitz.open(file_path)
        pages_text = []

        for page_num, page in enumerate(doc):
            # flags 参数控制提取行为：
            # TEXT_PRESERVE_WHITESPACE 保留空白，有助于保留表格结构
            text = page.get_text("text", sort=True)  # sort=True 按阅读顺序排序

            # 简单去除页眉页脚噪声：过滤掉极短的行（通常是页码）
            lines = text.split("\n")
            cleaned_lines = [l for l in lines if len(l.strip()) > 3]
            text = "\n".join(cleaned_lines)

            if text.strip():
                # 在每页文本开头插入页码标记
                # 这样切割后的 chunk 可以知道自己来自第几页
                pages_text.append(f"[第{page_num + 1}页]\n{text}")

        doc.close()
        return "\n\n".join(pages_text)

    def _parse_docx(self, file_path: str) -> str:
        """
        Word 文档的坑：
        1. 段落（paragraph）和表格（table）是两种不同的结构，要分别处理
        2. 表格里的文字如果只遍历 paragraphs 会漏掉
        3. 嵌套表格（表格中有表格）处理起来很麻烦，这里只处理一层
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = DocxDocument(file_path)
        content_parts = []

        for element in doc.element.body:
            # 判断元素类型
            tag = element.tag.split("}")[-1]  # 去掉命名空间前缀

            if tag == "p":  # 段落
                para_text = element.text_content() if hasattr(element, 'text_content') else ""
                # 更可靠的方式：从 paragraph 对象读
                pass
            elif tag == "tbl":  # 表格
                pass

        # 更简单可靠的方式：直接遍历 paragraph 和 table 对象
        content_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                content_parts.append(para.text)

        # 提取表格内容（转为文本）
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells
                )
                if row_text.strip(" |"):
                    table_rows.append(row_text)
            if table_rows:
                content_parts.append("\n".join(table_rows))

        return "\n\n".join(content_parts)

    def _infer_doc_type(self, filename: str) -> str:
        """
        根据文件名推断文档类型，用于元数据过滤。
        生产环境可以更精细，比如读文件头部内容来判断。
        """
        name = filename.lower()
        if any(k in name for k in ["hr", "人事", "员工", "假期", "薪酬"]):
            return "hr_policy"
        elif any(k in name for k in ["manual", "手册", "操作", "系统"]):
            return "it_manual"
        elif any(k in name for k in ["faq", "问答", "常见问题"]):
            return "faq"
        elif any(k in name for k in ["security", "安全", "信息安全"]):
            return "security_policy"
        else:
            return "general"


# ─────────────────────────────────────────
# 第二步：Chunking 策略
# ─────────────────────────────────────────

class RecursiveTextSplitter:
    """
    策略一：递归字符切割（最常用，生产首选）

    核心思想：
    按优先级尝试不同的分隔符，优先在"自然边界"切割：
      段落（\\n\\n）> 换行（\\n）> 句号 > 逗号 > 空格 > 单字符

    为什么"递归"：
    如果按 \\n\\n 切出来的块还是太大，就对这个块再用下一级分隔符切，
    直到满足 chunk_size 要求。

    chunk_overlap 的作用（非常重要）：
    问题：答案可能横跨两个 chunk 的边界。
    例如：
      chunk1: "...年假申请须提前3个工作日提交，"
      chunk2: "经直属主管审批后生效。当年未使用..."
    用户问"年假审批流程"，单独看 chunk1 或 chunk2 都不完整。

    解决：让相邻 chunk 有重叠部分（overlap），
    chunk1 末尾的内容也出现在 chunk2 开头，保证关键信息不会被切断。
    代价：存储量增加（overlap/chunk_size 的比例，通常 10%~20%）
    """

    def __init__(
        self,
        chunk_size: int = 500,      # 每个 chunk 的目标字符数
        chunk_overlap: int = 50,    # 相邻 chunk 的重叠字符数
        separators: Optional[list] = None,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        # 中文文档的分隔符优先级
        # 注意：中文没有空格分词，所以最后几个分隔符和英文不同
        self.separators = separators or [
            "\n\n",   # 段落分隔（最优先，语义最完整）
            "\n",     # 换行
            "。",     # 中文句号
            "；",     # 中文分号
            "，",     # 中文逗号
            ".",      # 英文句号
            " ",      # 空格
            "",       # 最后手段：强制按字符数切
        ]

    def split(self, document: Document) -> list[Chunk]:
        chunks_text = self._recursive_split(document.content, self.separators)
        chunks_with_overlap = self._add_overlap(chunks_text)

        chunks = []
        char_pos = 0  # 追踪每个 chunk 在原文中的起始位置

        for i, text in enumerate(chunks_with_overlap):
            # 尝试提取 chunk 所在的章节标题（用于元数据）
            section = self._extract_section(text, document.content)

            chunk_metadata = {
                **document.metadata,           # 继承文档级别的元数据
                "chunk_index": i,
                "chunk_total": len(chunks_with_overlap),
                "start_char":  char_pos,
                "char_count":  len(text),
                "section":     section,
                "content_hash": hashlib.md5(text.encode()).hexdigest(),
                "splitter":    "recursive",
            }

            chunks.append(Chunk(content=text, metadata=chunk_metadata))

            # 更新字符位置（减去 overlap 部分避免重复计数）
            char_pos += len(text) - self.chunk_overlap

        return chunks

    def _recursive_split(self, text: str, separators: list[str]) -> list[str]:
        """核心递归逻辑"""
        if not separators:
            # 没有分隔符了，强制按 chunk_size 切
            return [text[i:i+self.chunk_size]
                    for i in range(0, len(text), self.chunk_size)]

        separator = separators[0]
        splits = text.split(separator) if separator else list(text)

        result = []
        current = ""

        for split in splits:
            # 加上分隔符后看是否超长
            candidate = current + (separator if current else "") + split

            if len(candidate) <= self.chunk_size:
                current = candidate
            else:
                # 当前积累够了，先保存
                if current:
                    if len(current) > self.chunk_size:
                        # 单块就超长了，用下一级分隔符继续切
                        result.extend(
                            self._recursive_split(current, separators[1:])
                        )
                    else:
                        result.append(current)

                # 新块从当前 split 开始
                if len(split) > self.chunk_size:
                    result.extend(
                        self._recursive_split(split, separators[1:])
                    )
                    current = ""
                else:
                    current = split

        if current:
            if len(current) > self.chunk_size:
                result.extend(self._recursive_split(current, separators[1:]))
            else:
                result.append(current)

        # 过滤掉空块
        return [r for r in result if r.strip()]

    def _add_overlap(self, chunks: list[str]) -> list[str]:
        """
        在相邻 chunk 之间添加重叠。
        实现方式：每个 chunk（除第一个）的开头加上上一个 chunk 的末尾 overlap 字符。
        """
        if self.chunk_overlap == 0 or len(chunks) <= 1:
            return chunks

        result = [chunks[0]]
        for i in range(1, len(chunks)):
            # 取上一个 chunk 的末尾作为前缀
            prev_tail = chunks[i-1][-self.chunk_overlap:]
            result.append(prev_tail + chunks[i])

        return result

    def _extract_section(self, chunk_text: str, full_doc: str) -> str:
        """
        尝试找到 chunk 所属的章节标题。
        策略：在 chunk 文本之前，向上找最近的"第X章"或"X.X"格式标题。
        这是启发式方法，不能保证100%准确。
        """
        # 找到 chunk 在原文中的位置
        pos = full_doc.find(chunk_text[:50])  # 用前50字符定位
        if pos == -1:
            return ""

        # 向上搜索章节标题（正则匹配常见格式）
        preceding_text = full_doc[:pos]
        section_patterns = [
            r"第[一二三四五六七八九十\d]+章\s*\S+",  # 第X章 XXXX
            r"\d+\.\d+\s*\S+",                       # 1.1 XXXX
            r"[一二三四五六七八九十]+、\S+",           # 一、XXXX
        ]

        for pattern in section_patterns:
            matches = re.findall(pattern, preceding_text)
            if matches:
                return matches[-1][:30]  # 取最后一个（最近的），限制长度

        return ""


class SemanticChunker:
    """
    策略二：语义切割（精度更高，成本更高）

    核心思想：
    不按字符数切，而是按"语义边界"切。
    用 embedding 计算相邻句子的语义相似度，
    当相似度突然下降时，说明话题发生了转换，这里就是切割点。

    优点：chunk 内部语义更连贯，不会把一个完整的概念切断
    缺点：
      1. 需要对每个句子做 embedding，摄入速度慢（是递归切割的 10x 以上）
      2. embedding 有 API 调用成本（如果用 OpenAI）
      3. chunk 大小不均匀，可能产生极大或极小的 chunk

    适用场景：文档质量高、语义结构清晰、对召回精度要求很高的场景
    """

    def __init__(
        self,
        embedding_model,                  # sentence-transformers 模型
        breakpoint_threshold: float = 0.3, # 相似度下降超过此值 → 切割
        min_chunk_size: int = 100,
        max_chunk_size: int = 1000,
    ):
        self.model = embedding_model
        self.threshold = breakpoint_threshold
        self.min_size = min_chunk_size
        self.max_size = max_chunk_size

    def split(self, document: Document) -> list[Chunk]:
        # 第一步：按句子分割（比按字符更自然）
        sentences = self._split_to_sentences(document.content)

        if len(sentences) < 2:
            return [Chunk(content=document.content, metadata=document.metadata)]

        # 第二步：计算所有句子的 embedding
        # 注意：这里是批量计算，比逐句计算快很多
        embeddings = self.model.encode(sentences, batch_size=32, show_progress_bar=False)

        # 第三步：计算相邻句子的余弦相似度
        similarities = []
        for i in range(len(embeddings) - 1):
            sim = self._cosine_similarity(embeddings[i], embeddings[i+1])
            similarities.append(sim)

        # 第四步：找切割点
        # 切割点 = 相似度显著下降的位置（与前后平均值相比）
        breakpoints = self._find_breakpoints(similarities)

        # 第五步：按切割点合并句子成 chunk
        chunks_text = self._merge_sentences(sentences, breakpoints)

        # 处理过大/过小的 chunk
        chunks_text = self._normalize_chunk_sizes(chunks_text)

        # 构造 Chunk 对象
        chunks = []
        for i, text in enumerate(chunks_text):
            chunks.append(Chunk(
                content=text,
                metadata={
                    **document.metadata,
                    "chunk_index": i,
                    "chunk_total": len(chunks_text),
                    "splitter": "semantic",
                    "content_hash": hashlib.md5(text.encode()).hexdigest(),
                }
            ))

        return chunks

    def _split_to_sentences(self, text: str) -> list[str]:
        """
        中文断句。
        正则按句号、问号、感叹号切分，保留标点符号（不能丢）。
        """
        # 在句末标点后插入分隔符，再 split
        text = re.sub(r'([。！？\.\!\?])', r'\1<SPLIT>', text)
        sentences = text.split('<SPLIT>')
        # 过滤空句，合并过短的句子（少于10字的句子通常是标题或编号）
        result = []
        buffer = ""
        for s in sentences:
            s = s.strip()
            if not s:
                continue
            buffer += s
            if len(buffer) >= 20:  # 至少20字才算一个有意义的句子
                result.append(buffer)
                buffer = ""
        if buffer:
            result.append(buffer)
        return result

    def _cosine_similarity(self, a: np.ndarray, b: np.ndarray) -> float:
        """余弦相似度。embedding 已经 normalize 过则可直接用点积。"""
        return float(np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b) + 1e-8))

    def _find_breakpoints(self, similarities: list[float]) -> list[int]:
        """
        找语义边界。
        方法：计算每个位置与周围窗口平均相似度的差值，
        差值超过阈值 → 认为是话题转换点。
        """
        if not similarities:
            return []

        window = 3  # 用前后3个句子的平均值作为基准
        breakpoints = []

        for i, sim in enumerate(similarities):
            # 计算局部窗口平均
            start = max(0, i - window)
            end = min(len(similarities), i + window + 1)
            local_avg = np.mean(similarities[start:end])

            # 当前相似度比局部平均低 threshold → 切割点
            if local_avg - sim > self.threshold:
                breakpoints.append(i)

        return breakpoints

    def _merge_sentences(self, sentences: list[str], breakpoints: list[int]) -> list[str]:
        """按切割点把句子合并成 chunk"""
        if not breakpoints:
            return ["".join(sentences)]

        chunks = []
        start = 0
        for bp in breakpoints:
            chunk = "".join(sentences[start:bp+1])
            if chunk.strip():
                chunks.append(chunk)
            start = bp + 1

        # 最后一块
        last = "".join(sentences[start:])
        if last.strip():
            chunks.append(last)

        return chunks

    def _normalize_chunk_sizes(self, chunks: list[str]) -> list[str]:
        """
        处理异常大小的 chunk：
        - 太小（< min_size）：和下一个合并
        - 太大（> max_size）：用递归切割兜底
        """
        # 合并过小的 chunk
        merged = []
        buffer = ""
        for chunk in chunks:
            if len(buffer) + len(chunk) < self.min_size:
                buffer += chunk
            else:
                if buffer:
                    merged.append(buffer)
                buffer = chunk
        if buffer:
            merged.append(buffer)

        # 切割过大的 chunk（用递归切割兜底）
        result = []
        fallback_splitter = RecursiveTextSplitter(
            chunk_size=self.max_size,
            chunk_overlap=50
        )
        for chunk in merged:
            if len(chunk) > self.max_size:
                # 包装成临时 Document 用递归切割器处理
                tmp_doc = Document(content=chunk, metadata={})
                sub_chunks = fallback_splitter.split(tmp_doc)
                result.extend([c.content for c in sub_chunks])
            else:
                result.append(chunk)

        return result


# ─────────────────────────────────────────
# 第三步：摄入管道（把上面组合起来）
# ─────────────────────────────────────────

class IngestionPipeline:
    """
    把解析 + 切割 + 元数据 组合成一个完整的摄入流程。

    设计决策：
    - 默认用 RecursiveTextSplitter（快，效果够用）
    - SemanticChunker 作为可选项（精度高，成本高）
    - 支持增量摄入：通过 content_hash 跳过未变化的文档
    """

    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        use_semantic_chunking: bool = False,  # 是否用语义切割
        embedding_model=None,                  # 语义切割需要 embedding 模型
    ):
        self.parser = DocumentParser()

        if use_semantic_chunking and embedding_model:
            self.splitter = SemanticChunker(embedding_model)
            print("[IngestionPipeline] 使用语义切割（SemanticChunker）")
        else:
            self.splitter = RecursiveTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
            )
            print(f"[IngestionPipeline] 使用递归字符切割，chunk_size={chunk_size}, overlap={chunk_overlap}")

        self._processed_hashes: set[str] = set()  # 已处理文档的 hash，用于增量更新

    def ingest_file(self, file_path: str, extra_metadata: dict = None) -> list[Chunk]:
        """摄入单个文件，返回 chunk 列表"""
        print(f"\n[解析] {file_path}")

        # 解析
        doc = self.parser.parse(file_path)

        # 增量更新检查：hash 没变就跳过
        doc_hash = doc.metadata["content_hash"]
        if doc_hash in self._processed_hashes:
            print(f"  [跳过] 文档未变化（hash: {doc_hash[:8]}...）")
            return []
        self._processed_hashes.add(doc_hash)

        # 合并额外的元数据（调用方可以传入业务相关信息）
        if extra_metadata:
            doc.metadata.update(extra_metadata)

        # 切割
        chunks = self.splitter.split(doc)

        print(f"  [完成] 文档长度: {len(doc.content)} 字符 → {len(chunks)} 个 chunk")
        self._print_chunk_stats(chunks)

        return chunks

    def ingest_directory(self, dir_path: str, extra_metadata: dict = None) -> list[Chunk]:
        """摄入整个目录下的所有支持格式文件"""
        supported_ext = {".txt", ".md", ".pdf", ".docx"}
        all_chunks = []

        for file_path in Path(dir_path).rglob("*"):
            if file_path.suffix.lower() in supported_ext:
                try:
                    chunks = self.ingest_file(str(file_path), extra_metadata)
                    all_chunks.extend(chunks)
                except Exception as e:
                    print(f"  [错误] 处理 {file_path.name} 失败: {e}")

        print(f"\n[摄入完成] 共 {len(all_chunks)} 个 chunk")
        return all_chunks

    def _print_chunk_stats(self, chunks: list[Chunk]):
        """打印 chunk 统计信息，帮助调试切割效果"""
        if not chunks:
            return
        sizes = [len(c.content) for c in chunks]
        print(f"  chunk 统计: min={min(sizes)}, max={max(sizes)}, avg={sum(sizes)//len(sizes)}")
        # 如果最大 chunk 远超平均值，说明有段落没被正确切割，需要排查
        if max(sizes) > sum(sizes) / len(sizes) * 3:
            print(f"  ⚠️  警告: 最大 chunk 是平均值的 3 倍以上，请检查分隔符配置")


# ─────────────────────────────────────────
# 演示运行
# ─────────────────────────────────────────

if __name__ == "__main__":
    from rich.console import Console
    from rich.table import Table

    console = Console()

    data_dir = "./data"

    console.print("\n[bold cyan]═══ 文档摄入演示 ═══[/bold cyan]\n")

    # ── 演示1：递归字符切割（默认）
    console.print("[bold]1. 递归字符切割[/bold]")
    pipeline = IngestionPipeline(chunk_size=500, chunk_overlap=50)
    all_chunks = pipeline.ingest_directory(data_dir)

    # 展示前3个 chunk 的详情
    table = Table(title="前 3 个 Chunk 详情", show_lines=True)
    table.add_column("Chunk", style="cyan", width=8)
    table.add_column("内容（前80字）", width=50)
    table.add_column("来源", style="green")
    table.add_column("章节", style="yellow")
    table.add_column("字数")

    for chunk in all_chunks[:3]:
        table.add_row(
            f"#{chunk.metadata['chunk_index']}",
            chunk.content[:80].replace("\n", " ") + "...",
            chunk.metadata.get("source", ""),
            chunk.metadata.get("section", "")[:20],
            str(len(chunk.content)),
        )
    console.print(table)

    # ── 演示2：对比 overlap 的效果
    console.print("\n[bold]2. Overlap 效果对比[/bold]")
    splitter_no_overlap = RecursiveTextSplitter(chunk_size=200, chunk_overlap=0)
    splitter_with_overlap = RecursiveTextSplitter(chunk_size=200, chunk_overlap=50)

    test_doc = Document(
        content="年假申请须提前3个工作日通过OA系统提交。经直属主管审批后生效。当年未使用的年假可顺延至次年3月31日前使用，逾期作废，不予折现补偿。",
        metadata={"source": "test"}
    )

    chunks_no_ov = splitter_no_overlap.split(test_doc)
    chunks_with_ov = splitter_with_overlap.split(test_doc)

    console.print(f"无 overlap：{len(chunks_no_ov)} 个 chunk")
    for i, c in enumerate(chunks_no_ov):
        console.print(f"  [{i}] {c.content[:60]}")

    console.print(f"\n有 overlap（50字）：{len(chunks_with_ov)} 个 chunk")
    for i, c in enumerate(chunks_with_ov):
        console.print(f"  [{i}] {c.content[:60]}")

    console.print("\n[green]✓ 摄入层演示完成，下一步：02_indexing.py[/green]")
